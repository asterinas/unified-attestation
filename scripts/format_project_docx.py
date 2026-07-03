from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "项目说明文档.docx"
CHINESE_FONT = "仿宋"
FIXED_LINE_SPACING = Pt(25)


def set_style_font(style):
    if not hasattr(style, "font"):
        return
    from docx.oxml import OxmlElement

    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), CHINESE_FONT)


def set_run_chinese_font(run):
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        from docx.oxml import OxmlElement

        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), CHINESE_FONT)


def set_paragraph_line_spacing(paragraph):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = FIXED_LINE_SPACING
    fmt.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from iter_table_paragraphs(nested)


def main():
    doc = Document(str(DOCX))

    for style in doc.styles:
        set_style_font(style)
        if hasattr(style, "paragraph_format"):
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            style.paragraph_format.line_spacing = FIXED_LINE_SPACING
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        paragraphs.extend(iter_table_paragraphs(table))

    for paragraph in paragraphs:
        set_paragraph_line_spacing(paragraph)
        for run in paragraph.runs:
            set_run_chinese_font(run)

    doc.save(str(DOCX))
    print(DOCX)


if __name__ == "__main__":
    main()
