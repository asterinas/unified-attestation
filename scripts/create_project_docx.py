from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "项目说明文档.md"
OUTPUT = ROOT / "项目说明文档.docx"


def set_run_font(run, east_asia="Microsoft YaHei", ascii_font="Calibri"):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_text_with_inline_code(paragraph, text, bold=False):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, ascii_font="Consolas")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(31, 78, 121)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)
            run.bold = bold


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(24)
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(8)

    for name, size, color in (
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    if "CodeBlock" not in styles:
        code = styles.add_style("CodeBlock", 1)
        code.font.name = "Consolas"
        code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        code.font.size = Pt(9.5)
        code.paragraph_format.space_before = Pt(2)
        code.paragraph_format.space_after = Pt(2)
        code.paragraph_format.left_indent = Inches(0.12)


def add_cover(doc):
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Hydra 三方验证系统项目说明文档")
    set_run_font(run)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("安装步骤、运行方法与功能说明")
    set_run_font(run)
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(85, 85, 85)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("适用于 attester、verifier、relying-party 三方本地演示环境")
    set_run_font(run)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(85, 85, 85)

    doc.add_paragraph()
    box = doc.add_table(rows=1, cols=1)
    box.autofit = False
    cell = box.cell(0, 0)
    set_cell_shading(cell, "F2F4F7")
    set_cell_margins(cell, top=160, start=180, bottom=160, end=180)
    cell.text = ""
    p = cell.paragraphs[0]
    add_text_with_inline_code(
        p,
        "本文档说明项目的角色划分、环境安装、启动命令、passport/background_check 模式、"
        "本地数据存储、动态路径更新和常见问题处理。",
    )
    doc.add_section(WD_SECTION.NEW_PAGE)


def add_markdown_content(doc, markdown):
    in_code = False
    code_buffer = []

    def flush_code():
        nonlocal code_buffer
        if not code_buffer:
            return
        for line in code_buffer:
            p = doc.add_paragraph(style="CodeBlock")
            set_paragraph_shading(p, "F7F7F7")
            run = p.add_run(line if line else " ")
            set_run_font(run, east_asia="Microsoft YaHei", ascii_font="Consolas")
        code_buffer = []

    for raw in markdown.splitlines():
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
                code_buffer = []
            continue

        if in_code:
            code_buffer.append(line)
            continue

        if not line.strip():
            continue

        if line.startswith("# "):
            text = line[2:].strip()
            if text == "项目说明文档":
                continue
            p = doc.add_paragraph(style="Heading 1")
            add_text_with_inline_code(p, text, bold=True)
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_text_with_inline_code(p, line[3:].strip(), bold=True)
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_text_with_inline_code(p, line[4:].strip(), bold=True)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_text_with_inline_code(p, line[2:].strip())
        elif re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_text_with_inline_code(p, re.sub(r"^\d+\.\s+", "", line))
        else:
            p = doc.add_paragraph()
            add_text_with_inline_code(p, line)

    flush_code()


def main():
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_markdown_content(doc, markdown)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
